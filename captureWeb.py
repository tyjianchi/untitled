from urllib.parse import urlparse
import requests
import re
import json
import docx
import random
import string
from bs4 import BeautifulSoup  # 导入BeautifulSoup 模块

def generate_random_num(randomlength=13):
    str_list = [random.choice(string.digits) for i in range(randomlength)]
    random_str = ''.join(str_list)
    return random_str


def get_page(StartDate,EndtDate):
    url = 'http://rs.p5w.net/roadshow/getRoadshowList.shtml'
    pageList = {}
    formdata = {'perComType': 0, 'roadshowType': 5,'page':0,'start': StartDate, 'end': EndtDate}
    r = requests.post(url, data=formdata)
    r.raise_for_status()
    r.encoding = r.apparent_encoding
    row_cnt =json.loads(r.text)['total']
    if round(row_cnt/12)>=1:
        page_cnt=round(row_cnt/12)
    else:
        page_cnt=1

    for x in range(0,page_cnt):
        formdata_sub = {'perComType': 0, 'roadshowType': 5, 'page': x,'rows': 12, 'start': StartDate, 'end': EndtDate}
        r = requests.post(url, data=formdata_sub)
        r.raise_for_status()
        r.encoding = r.apparent_encoding
        p = json.loads(r.text)['rows']

        for i in p:
            if i['roadshowActiveHis']:
                pageList[i["roadshowTitle"]] = [i['roadshowActiveHis'], 0]
            else:
                pageList[i["roadshowTitle"]] = ['http://rs.p5w.net/' + i['roadshowUrl'], i['pid']]
    return pageList

def get_text(pageList):
    picList = []
    query_url='http://rs.p5w.net/roadshowLive/getNInteractionDatas.shtml'

    for title in pageList:
        pid=pageList[title][1]
        picList.append("活动标题:<" + title + ">")

        #根据不同传参方式请求地址
        if pid==0:
            now=generate_random_num()

            # 定义post传参
            bbs = '1'
            pageNo = '1'
            now_str = '1558155564709'
            flag_url=pageList[title][0].split('?')[-1].split('=')[-2]

            if re.match('.*selcode',flag_url):
                boardid = 'boardid=' + pageList[title][0].split('&')[-1].split('=')[-1]
                req_prefix = 'http://zsptbs.p5w.net/bbs/chatbbs/left.asp?&fsize=2'
                req_rul = req_prefix + '&pageNo='+pageNo+'&'+boardid
                # 请求
                req = requests.get(req_rul)
                req.raise_for_status()
                req.encoding = req.apparent_encoding

                # 格式化xml
                req_str = BeautifulSoup(req.text, 'lxml')
                req_table_all=req_str.find_all('table')[1]
                req_page_cnt=req_table_all.find_all('option')
                # 根据总页数循环处理
                for i in req_page_cnt:
                    r = requests.post(req_prefix + boardid + '&pageNo='+str(i.text)+'&'+boardid)
                    r.raise_for_status()
                    r.encoding = r.apparent_encoding
                    r_str = BeautifulSoup(r.text, 'lxml')
                    req_tab = r_str.find_all('table')[1]
                    req_t=req_tab.find_all('table')[1]
                    #headers = [c.get_text() for c in req_t.find('tr').find_all('td')[2:4]]
                    data = [[cell.get_text(strip=True).replace('\r\n','').replace('\u3000','') for cell in row.find_all('td')[2:4]]
                            for row in req_t.find_all("tr")]
                    for x in data[1:]:
                        for y in x:
                            picList.append(y)

            elif re.match('.*boardid',flag_url):
                boardid = 'boardid=' + pageList[title][0].split('?')[-1].split('=')[-1]
                req_prefix = 'http://newzspt.p5w.net/bbs/question_page.asp?'
                req_rul = req_prefix + boardid + '&bbs=' + bbs + '&pageNo=' + pageNo + '&now=' + now_str
                #请求
                req = requests.post(req_rul)
                req.raise_for_status()
                req.encoding = req.apparent_encoding

                #格式化xml
                req_str = BeautifulSoup(req.text, 'lxml')
                #根据总页数循环处理
                for i in range(int(req_str.q_page.text)):
                    r = requests.post(req_prefix+boardid+'&bbs='+bbs+'&pageNo='+str(i+1)+'&now='+now_str)
                    r.raise_for_status()
                    r.encoding = r.apparent_encoding
                    r_str = BeautifulSoup(r.text, 'lxml')
                    #picList.append("第"+str(i+1)+"页:" + r_str.text)
                    r_str1=[row.text.replace('\u3000','').replace('\r\n','').replace('&lt;br&gt;','') for row in r_str.find_all('q_and_r')]
                    for x in r_str1:
                        picList.append(x)
        else:
            formdata = {'roadshowId': pid}
            r = requests.post(query_url, data=formdata)
            r.raise_for_status()
            r.encoding = r.apparent_encoding
            p=json.loads(r.text)['rows']
            for i in p:
                if i['speakUserName']!="主持人":
                    #print("提 问 者:"+i['speakUserName'])
                    picList.append("提 问 者:"+i['speakUserName'])
                    #print("提问时间:"+i['speakTime'])
                    picList.append("提问时间:"+i['speakTime'])
                    if i['replyList']!=None:
                        #print("提问内容:"+i['speakContent'])
                        picList.append("提问内容:"+i['speakContent'])
                    else:
                        #print("提问内容:" + i['speakContent'])
                        str_ps_all = BeautifulSoup(i['speakContent'], 'lxml')
                        #print("提问内容:"+str_ps_all.text)
                        picList.append("提问内容:"+str_ps_all.text)

                    if i['replyList']!=None:
                        for x in i['replyList']:
                            if x['speakUserName']!=None:
                                #print("回 答 者:"+x['speakUserName'])
                                picList.append("回 答 者:"+x['speakUserName'])
                            if x['speakTime']!=None:
                                #print("回答时间:"+x['speakTime'])
                                picList.append("回答时间:"+x['speakTime'])
                            if x['speakContent']!=None:
                                #print("回答内容:"+x['speakContent'])
                                picList.append("回答内容:"+x['speakContent'])
    return picList

def writeDoc(picList):
    file = docx.Document()
    for i in picList:
        if re.match(r'活动标题',i):
           file.add_heading(i, level=1)
        else:
           file.add_paragraph(i)
    file.save(r"writeResult.docx")

if __name__ == '__main__':
    try:

        #设置开始时间
        StartDate='2010-05-08'
        #设置结束时间
        EndtDate='2010-05-30'

        #爬取每一页，找到每个说明会对应的url地址
        pageList = get_page(StartDate,EndtDate)

        #根据每个说明会对应的url地址爬取互动内容
        picList = get_text(pageList)

        #将最终的爬取结果在脚本的同目录下，写入writeResult.docx
        writeDoc(picList)

        print("爬取成功，请检查脚本同目录writeResult.docx文件")
    except Exception as e:
        print("爬取失败，有异常")
        print(e)